from flask import Flask, render_template, request, redirect
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase, Mapped, mapped_column
import docx


class Base(DeclarativeBase):
    pass


db = SQLAlchemy(model_class=Base)

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = "sqlite:///car_service.db"
db.init_app(app)


class Acceptence(db.Model):
    id: Mapped[int] = mapped_column(primary_key=True)
    name_of_work: Mapped[str] = mapped_column(nullable=False)
    cost: Mapped[int] = mapped_column(nullable=False)


class Stock(db.Model):
    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(nullable=False)
    cost: Mapped[int] = mapped_column(nullable=False)
    amount: Mapped[int] = mapped_column(nullable=False)
    brand: Mapped[str] = mapped_column(nullable=False)
    model: Mapped[str] = mapped_column(nullable=False)


class Clients(db.Model):
    id: Mapped[int] = mapped_column(primary_key=True)
    f: Mapped[str] = mapped_column(nullable=False)
    i: Mapped[str] = mapped_column(nullable=False)
    o: Mapped[str] = mapped_column(nullable=False)
    brand: Mapped[str] = mapped_column(nullable=False)
    model: Mapped[str] = mapped_column(nullable=False)


with app.app_context():
    db.create_all()


#главная
@app.route('/')
def home():
    return render_template('home.html')


#приемка
@app.route('/acceptence', methods=['POST', 'GET'])
def acceptence():
    work_show = Acceptence.query.order_by(Acceptence.name_of_work).all()
    if request.method == 'POST':
        works_ids = request.form.getlist('id')
        for work_id in works_ids:
            delete = db.session.query(Acceptence).get(int(work_id))
            db.session.delete(delete)
        db.session.commit()
        return redirect('/acceptence')
    return render_template('acceptence.html', work_show=work_show)


@app.route('/acceptence_add', methods=['POST', 'GET'])
def acceptence_add():
    if request.method == 'POST':
        acceptence_add = Acceptence(name_of_work=request.form['name'],
                                    cost=request.form['cost'],)
        db.session.add(acceptence_add)
        db.session.commit()
        return redirect('/acceptence_add')
    else:
        return render_template('acceptence_add.html')


@app.route('/acceptence/<int:id>/update', methods=['POST', 'GET'])
def acceptence_update(id):
    acceptence_update = Acceptence.query.get(id)
    if request.method == 'POST':
        acceptence_update.name_of_work = request.form['name']
        acceptence_update.cost = request.form['cost']
        db.session.commit()
        return redirect('/acceptence')
    else:
        return render_template('acceptence_update.html', acceptence_update=acceptence_update)


#склад
@app.route('/add_item', methods=['POST', 'GET'])
def add_item():
    if request.method == 'POST':
        stock_add = Stock(name=request.form['name'],
                          cost=request.form['cost'],
                          amount=request.form['amount'],
                          brand=request.form['brand'],
                          model=request.form['model'])
        db.session.add(stock_add)
        db.session.commit()
        return redirect('/add_item')
    else:
        return render_template('add_item.html')


@app.route('/stock', methods=['POST', 'GET'])
def Stock_show():
    stock_show = Stock.query.order_by(Stock.brand).all()
    if request.method == 'POST':
        items_ids = request.form.getlist('id')
        for item_id in items_ids:
            delete = db.session.query(Stock).get(int(item_id))
            db.session.delete(delete)
        db.session.commit()
        return redirect('/stock')
    return render_template('stock.html', stock_show=stock_show)


@app.route('/stock/<int:id>/update', methods=['POST', 'GET'])
def stock_update(id):
    stock_update = Stock.query.get(id)
    if request.method == 'POST':
        stock_update.name = request.form['name']
        stock_update.cost = request.form['cost']
        stock_update.amount = request.form['amount']
        stock_update.brand = request.form['brand']
        stock_update.model = request.form['model']
        db.session.commit()
        return redirect('/stock')
    else:
        return render_template('stock_update.html', stock_update=stock_update)


#клиенты
@app.route('/clients', methods=['POST', 'GET'])
def clients():
    clients_show = Clients.query.order_by(Clients.f).all()
    if request.method == 'POST':
        clients_ids = request.form.getlist('id')
        for clients_id in clients_ids:
            delete = db.session.query(Clients).get(int(clients_id))
            db.session.delete(delete)
        db.session.commit()
        return redirect('/clients')
    return render_template('clients.html', clients_show=clients_show)


@app.route('/clients_add', methods=['POST', 'GET'])
def clients_add():
    if request.method == 'POST':
        clients_add = Clients(f=request.form['f'],
                              i=request.form['i'],
                              o=request.form['o'],
                              brand=request.form['brand'],
                              model=request.form['model'])
        db.session.add(clients_add)
        db.session.commit()
        return redirect('/clients')
    else:
        return render_template('clients_add.html')


@app.route('/clients/<int:id>/update', methods=['POST', 'GET'])
def clients_update(id):
    clients_update = Clients.query.get(id)
    if request.method == 'POST':
        clients_update.f = request.form['f']
        clients_update.i = request.form['i']
        clients_update.o = request.form['o']
        clients_update.brand = request.form['brand']
        clients_update.model = request.form['model']
        db.session.commit()
        return redirect('/clients')
    else:
        return render_template('clients_update.html', clients_update=clients_update)


#печать
@app.route('/paste', methods=['POST', 'GET'])
def paste():
    acceptence_paste = Acceptence.query.order_by(Acceptence.name_of_work).all()
    stock_paste = Stock.query.order_by(Stock.name).all()
    if request.method == 'POST':
        items_id = request.form.getlist('id')
        doc = docx.Document()
        doc.add_paragraph('Проделанные работы')
        doc.add_paragraph('')
        amount_stock = 0
        amount_acceptence = 0
        for item_id in items_id:
            if 'a' in item_id:
                amount_stock += 1
            else:
                amount_acceptence += 1
        doc.add_paragraph('Работы')
        table_acceptence = doc.add_table(rows=amount_acceptence + 1, cols=3)
        table_acceptence.style = 'Table Grid'
        cell = table_acceptence.cell(0, 0)
        cell.text = 'Наименование'
        cell = table_acceptence.cell(0, 1)
        cell.text = 'Цена'
        cell = table_acceptence.cell(0, 2)
        cell.text = 'Кратность'
        doc.add_paragraph('')
        doc.add_paragraph('Детали')
        table_stock = doc.add_table(rows=amount_stock + 1, cols=6)
        table_stock.style = 'Table Grid'
        cell = table_stock.cell(0, 0)
        cell.text = 'Номер'
        cell = table_stock.cell(0, 1)
        cell.text = 'Наименование'
        cell = table_stock.cell(0, 2)
        cell.text = 'Цена'
        cell = table_stock.cell(0, 3)
        cell.text = 'Количество'
        cell = table_stock.cell(0, 4)
        cell.text = 'Марка'
        cell = table_stock.cell(0, 5)
        cell.text = 'Модель'
        i = 1
        n = 1
        for item_id in items_id:
            if 'a' in item_id:
                item_id = item_id.replace('a', '')
                stock_doc = Stock.query.get(int(item_id))
                cell = table_stock.cell(i, 0)
                cell.text = str(stock_doc.id)
                cell = table_stock.cell(i, 1)
                cell.text = stock_doc.name
                cell = table_stock.cell(i, 2)
                cell.text = str(int(stock_doc.cost) * int(request.form.getlist('amount_stock')[i - 1]))
                cell = table_stock.cell(i, 3)
                cell.text = str(request.form.getlist('amount_stock')[i - 1])
                cell = table_stock.cell(i, 4)
                cell.text = stock_doc.brand
                cell = table_stock.cell(i, 5)
                cell.text = stock_doc.model
                i += 1
            else:
                acceptence_doc = Acceptence.query.get(int(item_id))
                cell = table_acceptence.cell(n, 0)
                cell.text = acceptence_doc.name_of_work
                cell = table_acceptence.cell(n, 1)
                cell.text = str(acceptence_doc.cost * int(request.form.getlist('amount_acceptence')[n - 1]))
                cell = table_acceptence.cell(n, 2)
                cell.text = str(request.form.getlist('amount_acceptence')[n - 1])
                n += 1
        doc.save('template.docx')
    return render_template('paste.html', acceptence_paste=acceptence_paste, stock_paste=stock_paste)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
